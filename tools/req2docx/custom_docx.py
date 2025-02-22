"""Project specific docx converter functions.

    Author: Norbert Schulz (norbert.schulz@newtec.de)
"""

# pyTRLCConverter - A tool to convert TRLC files to specific formats.
# Copyright (c) 2024 - 2025 NewTec GmbH
#
# This file is part of pyTRLCConverter program.
#
# The pyTRLCConverter program is free software: you can redistribute it and/or modify it under
# the terms of the GNU General Public License as published by the Free Software Foundation,
# either version 3 of the License, or (at your option) any later version.
#
# The pyTRLCConverter program is distributed in the hope that it will be useful, but
# WITHOUT ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or
# FITNESS FOR A PARTICULAR PURPOSE. See the GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License along with pyTRLCConverter.
# If not, see <https://www.gnu.org/licenses/>.

# Imports **********************************************************************
import os
import docx
from pyTRLCConverter.docx_converter import DocxConverter
from pyTRLCConverter.ret import Ret
from pyTRLCConverter.plantuml import PlantUML
from pyTRLCConverter.trlc_helper import Record_Object

# Variables ********************************************************************

# Classes **********************************************************************

class CustomDocxConverter(DocxConverter):
    """Custom Project specific Docx format converter. 
    """
    def __init__(self, args: any) -> None:
        super().__init__(args)

        self._img_counter = 1

    @staticmethod
    def get_description() -> str:
        """ Return converter description.

         Returns:
            str: Converter description
        """
        return "Convert into project specific docx format."

    # pylint: disable=unused-argument
    def convert_record_object(self, record: Record_Object, level: int) -> Ret:
        """Converts a record object to docx format.

        Args:
            record (Record_Object): Record object to convert
            level (int): Current level of the record object
        
        Returns:
            Ret: Status
        """
        match record.n_typ.name:
            case "Info":
                ret = self._convert_record_object_info(record, level)
            case "Diagram":
                ret = self._convert_record_object_diagram(record, level)
            case _:
                ret = super().convert_record_object(record, level)

        return ret

    def _convert_record_object_info(self, record: Record_Object, level: int) -> Ret:
        """Convert an information record object to the destination format.

        Args:
            record (Record_Object): The record object to convert.
            level (int): Current level of the record object
        
        Returns:
            Ret: Status
        """
        attributes = record.to_python_dict()
        if "description" in attributes:
            self._docx.add_paragraph(self._get_attribute(record, "description"))

        return Ret.OK

    def _convert_record_object_diagram(self, record: Record_Object, level: int) -> Ret:
        """Convert a software diagram record object to the destination format.

        Args:
            record (Record_Object): The record object to convert.
            level (int): Current level of the record object
        
        Returns:
            Ret: Status
        """

        result = Ret.ERROR

        attributes = record.to_python_dict()
        file_path = self._locate_file(attributes.get("file_path"))
        if file_path is not None:
            puml = PlantUML()
            puml.generate("png", file_path, self._args.out)

            file_dst_path = os.path.basename(file_path)
            file_dst_path = os.path.splitext(file_dst_path)[0]
            file_dst_path += ".png"

            # PlantUML uses as output filename the diagram name if available.
            # The diagram name may differ from the filename.
            # To aovid that a invalid reference will be in the Markdown document,
            # ensure that the generated filename is as expected.
            expected_dst_path = os.path.join(self._args.out, file_dst_path)
            if os.path.isfile(expected_dst_path) is False:
                raise FileNotFoundError(
                    f"{file_path} diagram name ('@startuml <name>') may differ from file name,"
                    f"expected {expected_dst_path}."
                )
            p = self._docx.add_paragraph()
            p.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(expected_dst_path, width=docx.shared.Inches(6))
            run.add_text(f"Figure {self._img_counter} {self._get_attribute(record, 'caption')}")

            result = Ret.OK

        return result

# Functions ********************************************************************

# Main *************************************************************************
